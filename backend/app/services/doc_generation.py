"""
Document generation service.
Fills DOCX templates with customer data and converts to PDF.

PDF conversion uses LibreOffice headless on every platform (Windows, macOS,
Linux) so local output is byte-for-byte the same engine as production, where
the Docker image ships LibreOffice. There is intentionally no docx2pdf / MS
Word path — keeping a single conversion stack means "looks right locally"
also means "looks right in prod".
"""
import io
import os
import re
import shutil
import zipfile
import asyncio
import tempfile
import subprocess
from functools import lru_cache
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

from . import storage

TEMPLATE_DIR = Path(__file__).parent.parent.parent / "DOCS"

SIGNATURE_KEY = "signature.png"
PHOTO_KEY = "photo.jpg"
# Customer Aadhaar card images collected on the signing page.
AADHAR_FRONT_KEY = "aadhar_front.jpg"
AADHAR_BACK_KEY = "aadhar_back.jpg"

TEMPLATES = {
    "Annexure_1":           "TEMPELATE_Annexure.docx",
    "Aadhar":               "Aadhar.docx",
    "WCR":                  "WCR.docx",
    "Annexure_3":           "Annexure-3-Net-Metering.docx",
    "NP_Agreement":         "np_agreement.docx",
    "NP_Agreement_First_Page": "np_agreement_first_page.docx",
    "Meter_testing_letter": "Meter Testing Letter.docx",
    # DCR is no longer auto-generated — the admin uploads the official
    # government DCR PDF (see uploads router / DCR_UPLOAD_KEY below).
}

# ── Admin-uploaded documents (stored under a protected sub-prefix) ────────────
# These live under customers/{id}/uploads/ so the "delete all PDFs before
# regeneration" step never wipes them when the customer signs.
UPLOAD_PREFIX = "uploads/"
INSTALLATION_UPLOAD_KEY = UPLOAD_PREFIX + "installation.pdf"
DCR_UPLOAD_KEY          = UPLOAD_PREFIX + "dcr.pdf"
NP_STAMP_UPLOAD_KEY     = UPLOAD_PREFIX + "np_first_page_stamped.pdf"

# Friendly filenames for uploaded extras in the delivered bundle.
INSTALLATION_OUT_NAME = "Customer_Installation_Photo.pdf"
DCR_OUT_NAME          = "DCR.pdf"

# WCR and Aadhaar declaration are delivered as a single combined PDF (WCR pages
# first, then the Aadhaar declaration). They are filled and converted as separate
# templates, then merged into one file before upload so every downstream view
# (signing-page review, admin doc list, delivered bundle) shows just one document.
COMBINED_WCR_AADHAR_TYPE = "WCR_Aadhar"
WCR_DOC_TYPE = "WCR"
AADHAR_DOC_TYPE = "Aadhar"

# Internal artifacts that are never delivered to the customer.
HIDDEN_DOC_NAMES = {SIGNATURE_KEY, PHOTO_KEY, AADHAR_FRONT_KEY, AADHAR_BACK_KEY}
# Generated, unstamped NP first page — for the admin to print only.
NP_FIRST_PAGE_SUFFIX = "_NP_Agreement_First_Page.pdf"
# Base NP agreement (pages 2+). The stamped first page is merged in front of it.
NP_AGREEMENT_SUFFIX = "_NP_Agreement.pdf"

# placeholder -> (image-key, render width)
IMAGE_PLACEHOLDERS = {
    "${CUSTOMER_SIGN}$":   ("signature",    Inches(1.25)),
    "${CUSTOMER_PHOTO}$":  ("photo",        Inches(2.5)),
    # Aadhaar card images uploaded by the customer on the signing page.
    # Kept under half the ~6.27in usable page width so FRONT and BACK sit
    # side by side on one line (at 3.0in each they overflowed and wrapped).
    "${AADHAR_FRONT}$":    ("aadhar_front", Inches(2.8)),
    "${AADHAR_BACK}$":     ("aadhar_back",  Inches(2.8)),
}


def sanitize(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", str(name)).strip()


def build_replacements(customer: dict) -> dict:
    replacements = {}
    for key, value in customer.items():
        if key.startswith("_") or value is None:
            continue
        replacements[f"${{{key}}}$"] = str(value) if value else ""
    return replacements


# ── Placeholder discovery & pre-generation validation ─────────────────────────
# Templates use ${FIELD_NAME}$ placeholders. Before generating we scan every
# template, collect the data fields each one needs, and confirm the customer has
# a non-blank value for all of them — otherwise the produced documents would
# carry leftover ${...}$ markers or blanks. Templates are read only, never
# modified, here.

# Matches a single ${FIELD_NAME}$ placeholder and captures FIELD_NAME.
_PLACEHOLDER_RE = re.compile(r"\$\{([A-Za-z0-9_]+)\}\$")

# These placeholders are filled with images collected on the signing page (the
# customer's signature, photo and Aadhaar card scans), not from the customer
# data form, so they are intentionally excluded from the data-field check.
IMAGE_FIELD_KEYS = {k.strip("${}$") for k in IMAGE_PLACEHOLDERS}

# Documents that are generated only when all their required fields are present.
# These fields (meter make/number, receipt no., testing date) are filled in late
# — after the installation and meter testing are actually done — so leaving them
# blank must NOT block the rest of the (compulsory) documents. When they're
# missing we simply skip the optional document and generate everything else; the
# admin fills the meter details in later and regenerates to produce it.
OPTIONAL_DOCUMENTS = {"Meter_testing_letter"}

# Friendly, admin-facing labels for each generated document.
DOCUMENT_LABELS = {
    "Annexure_1":              "Annexure-1 — Solar Installation Declaration",
    "Aadhar":                  "Aadhaar Declaration",
    "WCR":                     "Work Completion Report",
    "Annexure_3":              "Annexure-3 — Net Metering",
    "NP_Agreement":            "Net-Metering Agreement",
    "NP_Agreement_First_Page": "Net-Metering Agreement (First Page)",
    "Meter_testing_letter":    "Meter Testing Letter",
}

# Friendly, admin-facing labels for each customer data field.
FIELD_LABELS = {
    "CONSUMER_NAME":            "Consumer Name",
    "CONSUMER_ADDRESS":         "Consumer Address",
    "CONSUMER_PHONE":           "Consumer Phone",
    "CONSUMER_EMAIL":           "Consumer Email",
    "CONSUMER_AADHAR":          "Consumer Aadhaar No.",
    "CONSUMER_NO":              "Consumer No.",
    "CONSUMER_APP_DATE":        "Application Date",
    "CONSUMER_APP_NO":          "Application No.",
    "DEALER_NAME":              "Dealer Name",
    "SANCTIONED_CAPACITY":      "Sanctioned Capacity",
    "SOLAR_CAPACITY":           "Solar Capacity",
    "INVERTER_MAKE":            "Inverter Make",
    "INVERTER_CAPACITY":        "Inverter Capacity",
    "INVERTER_GURANTEE":        "Inverter Guarantee",
    "INVERTER_SR_NO":           "Inverter Serial No.",
    "PANEL_COMPANY":            "Panel Company",
    "PANEL_WATT":               "Panel Wattage",
    "NO_OF_PANEL":              "Number of Panels",
    "TOTAL_PANEL_CAPACITY":     "Total Panel Capacity",
    "PANEL_SR_NO":              "Panel Serial No.",
    "PANEL_GURANTEE":           "Panel Guarantee",
    "INSTALLATION_DATE":        "Installation Date",
    "INSTALLATION_CITY":        "Installation City",
    "INSTALLATION_DISTRICT":    "Installation District",
    "DISCOM_REGISTERED_OFFICE": "DISCOM Registered Office",
    "SYSTEM_COST":              "System Cost",
    "METER_TESTING_DATE":       "Meter Testing Date",
    "METER_RECIPT_NO":          "Meter Receipt No.",
    "GENERATION_METER_MAKE":    "Generation Meter Make",
    "GENERATION_METER_NO":      "Generation Meter No.",
}


def field_label(key: str) -> str:
    """Human-readable label for a field key (falls back to a title-cased key)."""
    return FIELD_LABELS.get(key) or key.replace("_", " ").title()


def document_label(doc_type: str) -> str:
    """Human-readable label for a generated document type."""
    return DOCUMENT_LABELS.get(doc_type) or doc_type.replace("_", " ")


def _placeholders_in_template(template_path: Path) -> set[str]:
    """
    Return the set of field keys referenced by ${...}$ placeholders in a single
    template (paragraphs + table cells). Image placeholders are excluded — only
    data fields the admin fills in are returned.
    """
    doc = Document(template_path)
    keys: set[str] = set()

    def scan(paragraph):
        # Join runs so placeholders split across runs are still detected — this
        # mirrors how process_paragraph reconstructs the text before replacing.
        text = "".join(r.text for r in paragraph.runs)
        keys.update(_PLACEHOLDER_RE.findall(text))

    for para in doc.paragraphs:
        scan(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan(para)

    return keys - IMAGE_FIELD_KEYS


@lru_cache(maxsize=1)
def required_fields_by_document() -> dict[str, frozenset]:
    """
    Map each generated document type to the set of data fields it requires,
    discovered by scanning the templates. Cached because the templates don't
    change at runtime. Missing template files are skipped (they simply won't be
    generated either).
    """
    result: dict[str, frozenset] = {}
    for doc_type, template_file in TEMPLATES.items():
        template_path = TEMPLATE_DIR / template_file
        if not template_path.exists():
            continue
        result[doc_type] = frozenset(_placeholders_in_template(template_path))
    return result


def _has_value(customer: dict, key: str) -> bool:
    """True when the customer has a non-blank value for `key`."""
    value = customer.get(key)
    if value is None:
        return False
    return str(value).strip() != ""


def validate_customer_for_generation(customer: dict) -> dict | None:
    """
    Check that the customer has a value for every data field the templates need.

    Returns None when everything required is present. Otherwise returns a
    structured report the API/UI can show:
        {
          "missing_fields": [
            {"key": "SOLAR_CAPACITY", "label": "Solar Capacity",
             "documents": ["Work Completion Report", ...]},
            ...
          ],
          "documents": [
            {"document": "Work Completion Report",
             "missing": [{"key": "...", "label": "..."}], ...},
            ...
          ],
        }
    """
    required = required_fields_by_document()

    # field key -> document types that need it and are missing it
    missing_field_docs: dict[str, list[str]] = {}
    # document type -> missing field keys
    missing_by_doc: dict[str, list[str]] = {}

    for doc_type, fields in required.items():
        # Optional documents (e.g. Meter Testing Letter) never block generation —
        # their late-filled fields are checked separately and only cause that one
        # document to be skipped, not the whole compulsory set.
        if doc_type in OPTIONAL_DOCUMENTS:
            continue
        for key in fields:
            if not _has_value(customer, key):
                missing_field_docs.setdefault(key, []).append(doc_type)
                missing_by_doc.setdefault(doc_type, []).append(key)

    if not missing_field_docs:
        return None

    missing_fields = [
        {
            "key": key,
            "label": field_label(key),
            "documents": [document_label(dt) for dt in sorted(doc_types)],
        }
        for key, doc_types in sorted(missing_field_docs.items(), key=lambda kv: field_label(kv[0]))
    ]

    documents = [
        {
            "document": document_label(doc_type),
            "missing": [
                {"key": key, "label": field_label(key)}
                for key in sorted(keys, key=field_label)
            ],
        }
        for doc_type, keys in sorted(missing_by_doc.items(), key=lambda kv: document_label(kv[0]))
    ]

    return {"missing_fields": missing_fields, "documents": documents}


def skipped_optional_documents(customer: dict) -> list[dict]:
    """
    Optional documents that will NOT be generated because some of their required
    fields are still blank. Used to tell the admin (without blocking generation)
    which document — e.g. the Meter Testing Letter — was skipped and what to fill
    in to produce it on a later regeneration.

    Returns a list like:
        [{"document": "Meter Testing Letter",
          "missing": [{"key": "METER_TESTING_DATE", "label": "Meter Testing Date"}, ...]}]
    Empty when every optional document has all its fields.
    """
    required = required_fields_by_document()
    result: list[dict] = []
    for doc_type in sorted(OPTIONAL_DOCUMENTS, key=document_label):
        fields = required.get(doc_type)
        if not fields:
            continue
        missing = [key for key in fields if not _has_value(customer, key)]
        if missing:
            result.append({
                "document": document_label(doc_type),
                "missing": [
                    {"key": key, "label": field_label(key)}
                    for key in sorted(missing, key=field_label)
                ],
            })
    return result


def _copy_run_format(src_run, dst_run):
    """Copy character-level formatting from src_run to dst_run (font name, size, italic, underline, color)."""
    sf, df = src_run.font, dst_run.font
    df.name      = sf.name
    df.size      = sf.size
    df.italic    = sf.italic
    df.underline = sf.underline
    if sf.color and sf.color.type is not None:
        try:
            df.color.rgb = sf.color.rgb
        except Exception:
            pass


def process_paragraph(paragraph, replacements: dict, images: dict):
    all_keys = list(replacements.keys()) + list(IMAGE_PLACEHOLDERS.keys())
    original = "".join(r.text for r in paragraph.runs)
    if not any(k in original for k in all_keys):
        return

    # Capture formatting from the first non-empty run before clearing all runs.
    # All runs in a template paragraph are typically uniform, so this covers the
    # whole paragraph (font name, size, color, etc.).
    fmt_run = next((r for r in paragraph.runs if r.text), None) or (paragraph.runs[0] if paragraph.runs else None)

    for run in paragraph.runs:
        run.text = ""

    pattern = "|".join(re.escape(k) for k in all_keys)
    parts = re.split(f"({pattern})", original)
    for part in parts:
        if part in IMAGE_PLACEHOLDERS:
            img_key, width = IMAGE_PLACEHOLDERS[part]
            img_path = images.get(img_key)
            if img_path and img_path.exists():
                try:
                    new_run = paragraph.add_run()
                    if fmt_run:
                        _copy_run_format(fmt_run, new_run)
                    new_run.add_picture(str(img_path), width=width)
                except Exception as e:
                    print(f"Failed to embed image {img_path}: {e}")
            # else: leave blank — image not available yet
        elif part in replacements:
            new_run = paragraph.add_run(replacements[part])
            if fmt_run:
                _copy_run_format(fmt_run, new_run)
            new_run.bold = True  # filled values are always bolded
        else:
            new_run = paragraph.add_run(part)
            if fmt_run:
                _copy_run_format(fmt_run, new_run)


def fill_template(template_path: Path, output_path: Path, replacements: dict, images: dict):
    doc = Document(template_path)
    for para in doc.paragraphs:
        process_paragraph(para, replacements, images)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, replacements, images)
    doc.save(output_path)


@lru_cache(maxsize=1)
def _find_soffice() -> str:
    """
    Locate the LibreOffice executable. Honours the LIBREOFFICE_PATH env var,
    then falls back to PATH, then to the usual per-OS install locations.
    Raises RuntimeError with install guidance if LibreOffice is missing.
    """
    override = os.environ.get("LIBREOFFICE_PATH")
    if override and Path(override).exists():
        return override

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found

    candidates = [
        # Windows
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        # Linux (when not on PATH)
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
    ]
    for path in candidates:
        if Path(path).exists():
            return path

    raise RuntimeError(
        "LibreOffice not found. Install it so DOCX→PDF works the same as prod:\n"
        "  Windows: winget install --id TheDocumentFoundation.LibreOffice\n"
        "  macOS:   brew install --cask libreoffice\n"
        "  Linux:   apt-get install libreoffice\n"
        "Or set LIBREOFFICE_PATH to the soffice executable."
    )


def convert_docx_to_pdf(docx_path: Path) -> bool:
    """
    Convert a single DOCX to a PDF beside it using LibreOffice headless.

    Thin wrapper over convert_docx_batch — kept for single-file callers
    (e.g. the docgen_test.py harness). For generating a customer's full set,
    use convert_docx_batch so LibreOffice only starts up once.
    """
    produced = convert_docx_batch([docx_path], docx_path.parent)
    return docx_path.with_suffix(".pdf") in produced


def convert_docx_batch(docx_paths: list[Path], outdir: Path) -> set[Path]:
    """
    Convert many DOCX files to PDF in a single LibreOffice invocation.

    LibreOffice's ~1–2s startup is the expensive part, so converting a whole
    customer's documents in one `soffice` call instead of one-per-file is the
    dominant doc-gen speedup. If one document is malformed LibreOffice still
    converts the rest, so we report success per-file by checking which PDFs
    actually landed in `outdir`.

    A throwaway user-profile dir is passed via -env:UserInstallation so the
    conversion never collides with (or hangs behind) a LibreOffice window the
    user may already have open, and so concurrent batches stay isolated.

    Returns the set of PDF paths that were produced.
    """
    docx_paths = [p for p in docx_paths if p.exists()]
    if not docx_paths:
        return set()

    expected = {outdir / (p.stem + ".pdf") for p in docx_paths}
    profile_dir = Path(tempfile.mkdtemp(prefix="lo_profile_"))
    profile_uri = profile_dir.as_uri()
    # One LibreOffice startup serves the whole batch; scale the timeout with the
    # number of files (with generous headroom) rather than the old flat 120s.
    timeout = min(120 + 30 * len(docx_paths), 600)
    try:
        soffice = _find_soffice()
        result = subprocess.run(
            [soffice, "--headless", "--nologo", "--nofirststartwizard",
             f"-env:UserInstallation={profile_uri}",
             "--convert-to", "pdf", "--outdir", str(outdir),
             *[str(p) for p in docx_paths]],
            capture_output=True, text=True, timeout=timeout,
        )
        if result.returncode != 0:
            print(f"LibreOffice batch error: {result.stderr.strip() or result.stdout.strip()}")
    except subprocess.TimeoutExpired:
        print(f"PDF batch conversion timed out ({len(docx_paths)} files)")
    except Exception as e:
        print(f"PDF batch conversion failed: {e}")
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)

    produced = {p for p in expected if p.exists()}
    for missing in expected - produced:
        print(f"PDF failed: {missing.stem}")
    return produced


async def generate_for_customer(customer: dict) -> str:
    """
    Generate all documents for a customer in a temp dir, then upload the
    resulting PDFs to R2 under customers/{id}/. Returns the R2 prefix.
    """
    customer_id = str(customer.get("_id", "unknown"))
    safe_name = sanitize(customer.get("CONSUMER_NAME", "customer"))
    prefix = storage.customer_prefix(customer_id)

    replacements = build_replacements(customer)
    loop = asyncio.get_event_loop()

    work_dir = Path(tempfile.mkdtemp(prefix=f"docgen_{customer_id}_"))
    try:
        # Pull signature + Aadhaar images down from R2 (saved during signing) so
        # they get embedded. Fetch all four concurrently and skip the existence
        # HEAD — just attempt the download and treat a miss as "not present", so
        # each image costs one round trip instead of two serial ones.
        async def _fetch_image(img_key: str, fname: str):
            local = work_dir / fname
            try:
                data = await loop.run_in_executor(None, storage.download_bytes, prefix + fname)
            except Exception:
                return img_key, None
            local.write_bytes(data)
            return img_key, local

        fetched = await asyncio.gather(*(
            _fetch_image(k, f) for k, f in (
                ("signature", SIGNATURE_KEY), ("photo", PHOTO_KEY),
                ("aadhar_front", AADHAR_FRONT_KEY), ("aadhar_back", AADHAR_BACK_KEY),
            )
        ))
        images = {k: v for k, v in fetched}

        # Remove any stale PDFs from a previous generation before uploading fresh
        # ones — but never touch admin uploads under the uploads/ sub-prefix.
        # Delete them in one batched call rather than a LIST+DELETE per key.
        upload_prefix = prefix + UPLOAD_PREFIX
        stale = [o["key"] for o in await loop.run_in_executor(None, storage.list_objects, prefix)
                 if o["key"].lower().endswith(".pdf") and not o["key"].startswith(upload_prefix)]
        if stale:
            await loop.run_in_executor(None, storage.delete_keys, stale)

        # Phase 1 — fill every template's DOCX.
        required = required_fields_by_document()
        docx_paths: list[Path] = []
        for doc_type, template_file in TEMPLATES.items():
            template_path = TEMPLATE_DIR / template_file
            if not template_path.exists():
                print(f"Template not found: {template_path}")
                continue
            # Optional documents (e.g. Meter Testing Letter) are skipped when their
            # late-filled fields aren't ready yet, so the compulsory documents still
            # generate. The admin regenerates once the meter details are entered.
            if doc_type in OPTIONAL_DOCUMENTS:
                missing = [k for k in required.get(doc_type, ()) if not _has_value(customer, k)]
                if missing:
                    print(f"Skipping optional document {doc_type}; missing: {', '.join(missing)}")
                    continue
            docx_path = work_dir / f"{safe_name}_{doc_type}.docx"
            await loop.run_in_executor(None, fill_template, template_path, docx_path, replacements, images)
            docx_paths.append(docx_path)

        # Phase 2 — convert them all to PDF in a single LibreOffice startup.
        produced = await loop.run_in_executor(None, convert_docx_batch, docx_paths, work_dir)

        # Phase 2.5 — merge WCR + Aadhaar declaration into one combined PDF so the
        # customer (and admin download) get a single document instead of two.
        wcr_pdf = work_dir / f"{safe_name}_{WCR_DOC_TYPE}.pdf"
        aadhar_pdf = work_dir / f"{safe_name}_{AADHAR_DOC_TYPE}.pdf"
        if wcr_pdf in produced and aadhar_pdf in produced:
            combined_pdf = work_dir / f"{safe_name}_{COMBINED_WCR_AADHAR_TYPE}.pdf"
            await loop.run_in_executor(None, _merge_pdf_files, [wcr_pdf, aadhar_pdf], combined_pdf)
            produced = (produced - {wcr_pdf, aadhar_pdf}) | {combined_pdf}

        # Phase 3 — upload the produced PDFs to R2 in parallel (independent I/O).
        async def _upload(pdf_path: Path):
            await loop.run_in_executor(
                None, storage.upload_file, pdf_path, prefix + pdf_path.name, "application/pdf"
            )
            print(f"Uploaded to R2: {prefix + pdf_path.name}")

        await asyncio.gather(*(_upload(p) for p in sorted(produced)))
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    return prefix


# ── Upload presence (authoritative — reads R2, not the Mongo flags) ──────────

def uploaded_docs_present(prefix: str) -> dict:
    """Return which admin uploads actually exist in R2 under `prefix`."""
    return {
        "installation": storage.object_exists(prefix + INSTALLATION_UPLOAD_KEY),
        "np_stamp":     storage.object_exists(prefix + NP_STAMP_UPLOAD_KEY),
        "dcr":          storage.object_exists(prefix + DCR_UPLOAD_KEY),
    }


# ── Final-bundle assembly (delivered to the customer / admin download) ────────

def _merge_pdf_files(paths: list[Path], out_path: Path) -> None:
    """Concatenate the given PDF files (in order) into a single PDF at out_path."""
    out = fitz.open()
    opened = []
    try:
        for p in paths:
            doc = fitz.open(p)
            opened.append(doc)
            out.insert_pdf(doc)
        out.save(out_path)
    finally:
        out.close()
        for doc in opened:
            doc.close()


def merge_pdf_front(front: bytes, base: bytes) -> bytes:
    """Return a single PDF with `front` pages prepended to `base` pages."""
    out = fitz.open()
    front_doc = fitz.open(stream=front, filetype="pdf")
    base_doc = fitz.open(stream=base, filetype="pdf")
    try:
        out.insert_pdf(front_doc)
        out.insert_pdf(base_doc)
        return out.tobytes()
    finally:
        out.close()
        front_doc.close()
        base_doc.close()


def _is_deliverable(prefix: str, obj: dict) -> bool:
    """True for a generated PDF that should reach the customer."""
    key, name = obj["key"], obj["name"]
    if key.startswith(prefix + UPLOAD_PREFIX):      # admin uploads handled separately
        return False
    if name in HIDDEN_DOC_NAMES:                     # signature.png / photo.jpg
        return False
    if name.endswith(NP_FIRST_PAGE_SUFFIX):          # unstamped print copy
        return False
    return name.lower().endswith(".pdf")


# Friendly, customer-facing labels for the documents shown on the signing page.
# Keyed by the generated doc_type (the suffix of the R2 object name).
SIGNING_DOC_LABELS = {
    "Annexure_1":           "Annexure-1 — Solar Installation Declaration",
    "Aadhar":               "Aadhaar Declaration",
    "WCR":                  "Work Completion Report",
    "WCR_Aadhar":           "Work Completion Report & Aadhaar Declaration",
    "Annexure_3":           "Annexure-3 — Net Metering",
    "NP_Agreement":         "Net-Metering Agreement",
    "Meter_testing_letter": "Meter Testing Letter",
}


def signing_document_list(prefix: str, objects: list[dict], customer: dict) -> list[dict]:
    """
    Ordered list of documents shown to the customer on the signing page for
    review-before-consent. Mirrors `build_customer_bundle` so what they review
    is exactly what gets delivered (minus the signature/photo, which are added
    after they sign). Each item carries a stable `key`, a friendly `label`, and
    the internal fetch info (`r2_key`, `merge_stamp`) used by
    `fetch_signing_document`.
    """
    safe_name = sanitize(customer.get("CONSUMER_NAME", "customer"))
    keys = {o["key"]: o for o in objects}
    items: list[dict] = []

    # Generated deliverables, sorted by name for a stable, deterministic order.
    for o in sorted((o for o in objects if _is_deliverable(prefix, o)), key=lambda o: o["name"]):
        name = o["name"]
        doc_type = name[len(safe_name) + 1:-4] if name.startswith(safe_name + "_") else name[:-4]
        label = SIGNING_DOC_LABELS.get(doc_type) or doc_type.replace("_", " ")
        items.append({
            "label": label,
            "r2_key": o["key"],
            "merge_stamp": name.endswith(NP_AGREEMENT_SUFFIX),
        })

    # Admin uploads delivered alongside the generated docs.
    if prefix + INSTALLATION_UPLOAD_KEY in keys:
        items.append({"label": "Installation Photo", "r2_key": prefix + INSTALLATION_UPLOAD_KEY, "merge_stamp": False})
    if prefix + DCR_UPLOAD_KEY in keys:
        items.append({"label": "DCR Declaration", "r2_key": prefix + DCR_UPLOAD_KEY, "merge_stamp": False})

    for i, item in enumerate(items):
        item["key"] = str(i)
    return items


async def build_signing_manifest(prefix: str, customer: dict) -> list[dict]:
    """
    Build the signing-page document manifest (the same list `signing_document_list`
    produces) by reading R2 once. Persisted on the customer at send-link time so
    each review-page document view resolves from the DB instead of re-LISTing the
    whole R2 folder every time.
    """
    loop = asyncio.get_event_loop()
    objects = await loop.run_in_executor(None, storage.list_objects, prefix)
    return signing_document_list(prefix, objects, customer)


async def resolve_signing_item(prefix: str, customer: dict, key: str,
                               manifest: list[dict] | None = None) -> dict | None:
    """
    Resolve a signing-page document `key` to its manifest item
    ({label, r2_key, merge_stamp, key}) without downloading anything. Uses the
    persisted `manifest` when given (no R2 LIST); otherwise rebuilds it from R2
    for links sent before the manifest existed. Returns the item or None.
    """
    if manifest is None:
        loop = asyncio.get_event_loop()
        objects = await loop.run_in_executor(None, storage.list_objects, prefix)
        manifest = signing_document_list(prefix, objects, customer)
    return next((it for it in manifest if it["key"] == key), None)


async def fetch_signing_document(prefix: str, customer: dict, key: str,
                                 manifest: list[dict] | None = None) -> tuple[str, bytes] | None:
    """
    Fetch a single signing-page document's PDF bytes by its stable `key`.
    Applies the same stamped-first-page merge as the delivered bundle so the
    NP Agreement the customer reviews matches the final document. Returns
    (label, pdf_bytes) or None if the key doesn't resolve.
    """
    item = await resolve_signing_item(prefix, customer, key, manifest)
    if not item:
        return None

    loop = asyncio.get_event_loop()
    data = await loop.run_in_executor(None, storage.download_bytes, item["r2_key"])
    if item["merge_stamp"]:
        stamp_key = prefix + NP_STAMP_UPLOAD_KEY
        if await loop.run_in_executor(None, storage.object_exists, stamp_key):
            stamp_bytes = await loop.run_in_executor(None, storage.download_bytes, stamp_key)
            data = merge_pdf_front(stamp_bytes, data)
    return item["label"], data


def list_customer_documents(prefix: str, objects: list[dict]) -> list[dict]:
    """Logical document list (name/size) for the admin panel — no downloads."""
    docs = [
        {"name": o["name"], "size": o["size"], "type": "pdf"}
        for o in objects if _is_deliverable(prefix, o)
    ]
    by_key = {o["key"]: o for o in objects}
    if prefix + INSTALLATION_UPLOAD_KEY in by_key:
        docs.append({"name": INSTALLATION_OUT_NAME, "size": by_key[prefix + INSTALLATION_UPLOAD_KEY]["size"], "type": "pdf"})
    if prefix + DCR_UPLOAD_KEY in by_key:
        docs.append({"name": DCR_OUT_NAME, "size": by_key[prefix + DCR_UPLOAD_KEY]["size"], "type": "pdf"})
    return sorted(docs, key=lambda d: d["name"])


async def build_customer_bundle(prefix: str) -> list[tuple[str, bytes]]:
    """
    Assemble the final set of PDFs delivered to the customer:
      • generated docs (Annexure, Aadhar, WCR, etc.)
      • NP Agreement with the admin's stamped first page merged in front
      • admin-uploaded installation photo PDF and government DCR PDF
    Internal artifacts (signature, photo, unstamped first page) are excluded.
    """
    loop = asyncio.get_event_loop()
    objects = await loop.run_in_executor(None, storage.list_objects, prefix)
    keys = {o["key"] for o in objects}

    stamp_key = prefix + NP_STAMP_UPLOAD_KEY
    stamp_bytes = None
    if stamp_key in keys:
        stamp_bytes = await loop.run_in_executor(None, storage.download_bytes, stamp_key)

    # Plan the bundle first — (output name, R2 key, needs stamp merge) — in a
    # stable order, then download every file concurrently and assemble.
    plan: list[tuple[str, str, bool]] = []
    for o in sorted(objects, key=lambda o: o["name"]):
        if not _is_deliverable(prefix, o):
            continue
        plan.append((o["name"], o["key"], o["name"].endswith(NP_AGREEMENT_SUFFIX)))
    for upload_key, out_name in ((prefix + INSTALLATION_UPLOAD_KEY, INSTALLATION_OUT_NAME),
                                 (prefix + DCR_UPLOAD_KEY, DCR_OUT_NAME)):
        if upload_key in keys:
            plan.append((out_name, upload_key, False))

    async def _fetch(out_name: str, r2_key: str, merge: bool) -> tuple[str, bytes]:
        data = await loop.run_in_executor(None, storage.download_bytes, r2_key)
        if merge and stamp_bytes:
            data = merge_pdf_front(stamp_bytes, data)
        return out_name, data

    # gather preserves input order, so the bundle stays deterministic.
    return list(await asyncio.gather(*(_fetch(n, k, m) for n, k, m in plan)))


def zip_bundle(bundle: list[tuple[str, bytes]]) -> io.BytesIO:
    """
    Pack a list of (filename, bytes) into an in-memory ZIP.

    Uses ZIP_STORED (no compression): the bundle is entirely PDFs, which are
    already compressed, so DEFLATE would burn CPU for ~no size reduction. This
    is synchronous CPU work — call it via asyncio.to_thread from async handlers
    so it never blocks the event loop.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as zf:
        for name, data in bundle:
            zf.writestr(name, data)
    buf.seek(0)
    return buf
