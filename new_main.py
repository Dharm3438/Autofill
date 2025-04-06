from docx import Document
import re
import pandas as pd
import os
from pathlib import Path

def sanitize_filename(name):
    """Sanitize strings to be safe for filenames"""
    return re.sub(r'[\\/*?:"<>|]', "_", str(name)).strip()

# def create_replacements_dict(row):
#     """Create replacement dictionary from Excel row data"""
#     return {
#         f"${{{key}}}$": str(value) if pd.notna(value) else ""
#         for key, value in row.items()
#     }

def create_replacements_dict(row):
    """Create replacement dictionary from Excel row data with date formatting"""
    replacements = {}
    date_fields = {'CONSUMER_APP_DATE', 'INSTALLATION_DATE'}  # Add all date field names here
    
    for key, value in row.items():
        if key in date_fields and pd.notna(value):
            # Convert datetime to DD-MM-YYYY format
            if hasattr(value, 'strftime'):  # Check if it's a datetime object
                formatted_date = value.strftime('%d-%m-%Y')
            else:
                # Fallback for string dates
                try:
                    formatted_date = pd.to_datetime(value).strftime('%d-%m-%Y')
                except:
                    formatted_date = str(value)
            replacements[f"${{{key}}}$"] = formatted_date
        else:
            replacements[f"${{{key}}}$"] = str(value) if pd.notna(value) else ""
    
    return replacements

def process_paragraph(paragraph, replacements):
    """Process a single paragraph with replacements"""
    original_text = ''.join(run.text for run in paragraph.runs)
    if not any(key in original_text for key in replacements):
        return
        
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''
        
    # Rebuild with replacements
    parts = re.split(f"({'|'.join(re.escape(k) for k in replacements)})", original_text)
    for part in parts:
        new_run = paragraph.add_run(part if part not in replacements else replacements[part])
        if part in replacements:
            new_run.bold = True
            new_run.underline = True

def process_template(template_path, output_path, replacements):
    """Process a single template file with replacements"""
    doc = Document(template_path)
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, replacements)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, replacements)
    
    doc.save(output_path)

def generate_documents():
    # Configuration
    INPUT_EXCEL = "input_data.xlsx"
    TEMPLATE_DIR = "DOCS"
    OUTPUT_DIR = "generated_docs"
    
    # Template files to process (key: display name, value: filename)
    TEMPLATES = {
        "DCR": "DCR.docx",
        "Annexure-1": "TEMPELATE_Annexure.docx",
        "Aadhar": "Aadhar.docx",
        "WCR": "WCR.docx",
        "Annexure-3": "Annexure-3-Net-Metering.docx",
        "NP-Agreement": "np_agreement.docx",
    }

    # Read input data
    df = pd.read_excel(INPUT_EXCEL)
    
    # Create output directory
    Path(OUTPUT_DIR).mkdir(exist_ok=True)
    
    # Process each consumer
    for idx, row in df.iterrows():
        consumer_id = str(row.get("CONSUMER_APP_NO", f"consumer_{idx}"))
        consumer_name = str(row.get("CONSUMER_NAME", "consumer"))
        
        # Sanitize names for filesystem safety
        safe_consumer_id = sanitize_filename(consumer_id)
        safe_consumer_name = sanitize_filename(consumer_name)
        
        # Create consumer directory
        consumer_dir = Path(OUTPUT_DIR) / safe_consumer_name
        consumer_dir.mkdir(exist_ok=True)
        
        replacements = create_replacements_dict(row)
        print(replacements)
        # Generate all documents
        for doc_type, template_file in TEMPLATES.items():
            template_path = Path(TEMPLATE_DIR) / template_file
            output_filename = f"{safe_consumer_name}_{doc_type}.docx"
            output_path = consumer_dir / output_filename
            
            try:
                process_template(template_path, output_path, replacements)
                print(f"Generated {doc_type} for {safe_consumer_name}")
            except Exception as e:
                print(f"Error generating {doc_type} for {safe_consumer_name}: {str(e)}")
        

if __name__ == "__main__":
    generate_documents()
    print("Document generation completed!")