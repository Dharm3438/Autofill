from docx import Document
import re
import pandas as pd

def replace_text_in_paragraph(paragraph, replacements):
    # Combine all run texts into a single string
    full_text = ''.join(run.text for run in paragraph.runs)

    # Create a pattern to match all replacement keys
    pattern = '|'.join(re.escape(k) for k in replacements.keys())
    parts = re.split(f'({pattern})', full_text)

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''

    # Rebuild paragraph with styled replacements
    for part in parts:
        run = paragraph.add_run()
        if part in replacements:
            run.text = replacements[part]
            run.bold = True
            run.underline = True
        else:
            run.text = part


def replace_text_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

def fill_template(template_path, output_path, replacements):
    doc = Document(template_path)

    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, replacements)

    doc.save(output_path)

# Example usage
replacements = {
    "${SOLAR_CAPACITY}$": "3",
    "${CONSUMER_NAME}$": "Indresh Devji Poladiya",
    "${CONSUMER_ADDRESS}$": "Ganesh Cross Road, Near Yogeshwar Medical Chalisgaon",
    "${CONSUMER_APP_NO}$": "123456789011",
    "${CONSUMER_APP_DATE}$": "01/01/2025",
    "${TOTAL_PANEL_CAPACITY}$": "3150 Watt",
    "${NO_OF_PANEL}$": "6",
    "${PANEL_SR_NO}$": "5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757",
    "${PANEL_COMPANY}$": "Fujiyama Power System Private Limited",
    "${CELL_MANUFACTURER}$": "Mundra Solar Private Limited",
    "${PANEL_WATT}$": "525",
    "${CONSUMER_NO}$": "123456789",
    "${CONSUMER_PHONE}$": "7588647905",
    "${CONSUMER_EMAIL}$": "dharmesh@gmail.com",
    "${SANCTIONED_CAPACITY}$": "3",
    "${INSTALLATION_DATE}$": "01/01/2025",
    "${INVERTER_CAPACITY}$": "5",
    "${INVERTER_MAKE}$": "UTL",
    "${INSTALLATION_CITY}$": "Chalisgaon",
    "${INSTALLATION_DISTRICT}$": "Jalgaon",
    "${CONSUMER_AADHAR}$": "1234567890",
    "${SANCTION_NUMBER}$": "785267856725",
    "${PANEL_GURANTEE}$": "25 Years",
    "${INVERTER_GURANTEE}$": "21 Years",

    "${DISCOM_REGISTERED_OFFICE}$": "Chalisgaon-U",
}

fill_template("DOCS\DCR.docx", "res_DCR.docx", replacements)
fill_template("DOCS\TEMPELATE_Annexure.docx", "res_Annexure.docx", replacements)
fill_template("DOCS\Aadhar.docx", "res_Aadhar.docx", replacements)
fill_template("DOCS\WCR.docx", "res_WCR.docx", replacements)
# fill_template("DOCS\Annexure-3-Net-Metering.docx", "res_Annexure.docx", replacements)