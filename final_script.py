from docx import Document
from docx.shared import Pt
import re
import pandas as pd
import os
from pathlib import Path
from docx2pdf import convert
from PyPDF2 import PdfMerger
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

MAX_PDF_SIZE_MB = 0.3

def sanitize_filename(name):
    """Sanitize strings to be safe for filenames"""
    return re.sub(r'[\\/*?:"<>|]', "_", str(name)).strip()

def convert_to_pdf(docx_path, pdf_path, max_size_mb=10):
    """Convert DOCX to PDF with size limit check"""
    try:
        # Convert to PDF
        convert(docx_path, pdf_path)
        
        # Check file size
        pdf_size = os.path.getsize(pdf_path) / (1024 * 1024)  # Size in MB
        if pdf_size > max_size_mb:
            print(f"Warning: PDF {pdf_path} exceeds {max_size_mb}MB ({pdf_size:.2f}MB)")
            return False
        return True
    except Exception as e:
        print(f"PDF conversion failed for {docx_path}: {str(e)}")
        return False

def combine_pdf_files(input_paths, output_path):
    """Combine multiple PDF files into one"""
    try:
        merger = PdfMerger()
        for pdf_path in input_paths:
            merger.append(pdf_path)
        merger.write(output_path)
        merger.close()
        return True
    except Exception as e:
        print(f"Error combining PDF files: {str(e)}")
        return False

def create_replacements_dict(row):
    """Create replacement dictionary from Excel row data with date formatting"""
    replacements = {}
    date_fields = {'CONSUMER_APP_DATE', 'INSTALLATION_DATE', 'METER_TESTING_DATE'}  # Add all date field names here
    
    for key, value in row.items():
        if key in date_fields and pd.notna(value):
            # Convert datetime to DD-MM-YYYY format
            if hasattr(value, 'strftime'):  # Check if it's a datetime object
                formatted_date = value.strftime('%d-%m-%Y')
            else:
                # Fallback for string dates
                try:
                    formatted_date = pd.to_datetime(value).strftime('%d-%m-%Y')
                except:
                    formatted_date = str(value)
            replacements[f"${{{key}}}$"] = formatted_date
        else:
            replacements[f"${{{key}}}$"] = str(value) if pd.notna(value) else ""
    
    return replacements

def process_paragraph(paragraph, replacements):
    """Process a single paragraph with replacements"""
    original_text = ''.join(run.text for run in paragraph.runs)
    if not any(key in original_text for key in replacements):
        return
        
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''
        
    # Rebuild with replacements
    parts = re.split(f"({'|'.join(re.escape(k) for k in replacements)})", original_text)
    for part in parts:
        new_run = paragraph.add_run(part if part not in replacements else replacements[part])
        if part in replacements:
            new_run.bold = True
            # new_run.underline = True

def process_template(template_path, output_path, replacements):
    """Process a single template file with replacements"""
    doc = Document(template_path)
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, replacements)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, replacements)
    
    doc.save(output_path)

def process_consumer(idx, row, templates, template_dir, output_dir, dir_lock):
    """Process a single consumer's documents (called by worker threads)"""
    consumer_id = str(row.get("CONSUMER_APP_NO", f"consumer_{idx}"))
    consumer_name = str(row.get("CONSUMER_NAME", "consumer"))
    
    # Sanitize names for filesystem safety
    safe_consumer_id = sanitize_filename(consumer_id)
    safe_consumer_name = sanitize_filename(consumer_name)
    
    # Create consumer directory (with thread-safe lock)
    with dir_lock:
        consumer_dir = Path(output_dir) / safe_consumer_name
        
        # Skip if directory already exists
        if consumer_dir.exists():
            print(f"Consumer directory '{safe_consumer_name}' already exists. Skipping...")
            return
        
        consumer_dir.mkdir(exist_ok=True)
    
    replacements = create_replacements_dict(row)
    
    # Create panel serial numbers Excel file
    serial_numbers_string = replacements['${PANEL_SR_NO}$']
    serial_number_list = [serial_number.strip() for serial_number in serial_numbers_string.split(',')]
    df = pd.DataFrame(serial_number_list)
    excel_file_name = consumer_dir / f"{safe_consumer_name}_panel_serial_numbers.xlsx"
    df.to_excel(excel_file_name, index=False, header=False)
    print(f"Excel file '{excel_file_name}' has been created successfully.")

    # Process each template
    for doc_type, template_file in templates.items():
        template_path = Path(template_dir) / template_file
        base_filename = f"{safe_consumer_name}_{doc_type}"

        # Generate DOCX
        docx_path = consumer_dir / f"{base_filename}.docx"
        try:
            process_template(template_path, docx_path, replacements)
            print(f"Generated DOCX {doc_type} for {safe_consumer_name}")

            # Generate PDF
            pdf_path = consumer_dir / f"{base_filename}.pdf"
            if convert_to_pdf(docx_path, pdf_path, MAX_PDF_SIZE_MB):
                print(f"Generated PDF {doc_type} for {safe_consumer_name}")
            else:
                print(f"PDF generation failed for {doc_type} ({safe_consumer_name})")

        except Exception as e:
            print(f"Error generating {doc_type} for {safe_consumer_name}: {str(e)}")
    
    # Combine WCR and Aadhar documents
    wcr_pdf = consumer_dir / f"{safe_consumer_name}_WCR.pdf"
    aadhar_pdf = consumer_dir / f"{safe_consumer_name}_Aadhar.pdf"
    combined_pdf = consumer_dir / f"{safe_consumer_name}_WCR_Aadhar_Combined.pdf"
    
    if wcr_pdf.exists() and aadhar_pdf.exists():
        if combine_pdf_files([wcr_pdf, aadhar_pdf], combined_pdf):
            print(f"Generated Combined PDF for {safe_consumer_name}")
        else:
            print(f"Combined PDF generation failed for {safe_consumer_name}")
    else:
        print(f"Could not find WCR or Aadhar documents for {safe_consumer_name}")

def generate_documents(max_workers=4):
    """Generate documents for all consumers using parallel threading"""
    # Configuration
    INPUT_EXCEL = "input_data.xlsx"
    TEMPLATE_DIR = "DOCS"
    OUTPUT_DIR = "generated_docs"
    
    # Template files to process (key: display name, value: filename)
    TEMPLATES = {
        "Annexure_1": "TEMPELATE_Annexure.docx",
        "Aadhar": "Aadhar.docx",
        "WCR": "WCR.docx",
        "Annexure_3": "Annexure-3-Net-Metering.docx",
        "NP_Agreement": "np_agreement.docx",
        "Meter_testing_letter": "Meter Testing Letter.docx",
        "DCR": "DCR.docx"
    }

    # Read input data
    df = pd.read_excel(INPUT_EXCEL)
    
    # Create output directory
    Path(OUTPUT_DIR).mkdir(exist_ok=True)
    
    # Lock for thread-safe directory creation
    dir_lock = threading.Lock()
    
    print(f"Starting document generation for {len(df)} consumers using {max_workers} worker threads...")
    
    # Use ThreadPoolExecutor for parallel processing
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = []
        
        # Submit all consumer processing tasks
        for idx, row in df.iterrows():
            future = executor.submit(
                process_consumer, 
                idx, 
                row, 
                TEMPLATES, 
                TEMPLATE_DIR, 
                OUTPUT_DIR, 
                dir_lock
            )
            futures.append(future)
        
        # Wait for all tasks to complete and handle any exceptions
        completed = 0
        for future in as_completed(futures):
            try:
                future.result()
                completed += 1
            except Exception as e:
                print(f"Error in consumer processing: {str(e)}")

if __name__ == "__main__":
    generate_documents(max_workers=4)  # Adjust max_workers based on your CPU cores (2-8 recommended)
    print("Document generation completed!")