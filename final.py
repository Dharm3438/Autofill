from docx import Document
from docx.shared import Pt
import re
import pandas as pd
import os
from pathlib import Path
from docx2pdf import convert
from PyPDF2 import PdfMerger
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import time

MAX_PDF_SIZE_MB = 0.3

# ──────────────────────────────────────────────
# Shared state for tracking PDF conversion queue
# ──────────────────────────────────────────────
pdf_queue = []          # List of (docx_path, pdf_path) tuples
pdf_queue_lock = threading.Lock()


def sanitize_filename(name):
    """Sanitize strings to be safe for filenames."""
    return re.sub(r'[\\/*?:"<>|]', "_", str(name)).strip()


def convert_to_pdf(docx_path, pdf_path, max_size_mb=10):
    """Convert a single DOCX to PDF with optional size-limit check."""
    try:
        convert(docx_path, pdf_path)
        pdf_size = os.path.getsize(pdf_path) / (1024 * 1024)
        if pdf_size > max_size_mb:
            print(f"  ⚠  PDF exceeds {max_size_mb} MB ({pdf_size:.2f} MB): {pdf_path}")
            return False
        return True
    except Exception as e:
        print(f"  ✗  PDF conversion failed [{docx_path}]: {e}")
        return False


def combine_pdf_files(input_paths, output_path):
    """Merge multiple PDFs into one."""
    try:
        merger = PdfMerger()
        for p in input_paths:
            merger.append(p)
        merger.write(output_path)
        merger.close()
        return True
    except Exception as e:
        print(f"  ✗  PDF merge failed: {e}")
        return False


def create_replacements_dict(row):
    """Build placeholder → value map from an Excel row."""
    replacements = {}
    date_fields = {"CONSUMER_APP_DATE", "INSTALLATION_DATE", "METER_TESTING_DATE"}

    for key, value in row.items():
        if key in date_fields and pd.notna(value):
            if hasattr(value, "strftime"):
                formatted = value.strftime("%d-%m-%Y")
            else:
                try:
                    formatted = pd.to_datetime(value).strftime("%d-%m-%Y")
                except Exception:
                    formatted = str(value)
            replacements[f"${{{key}}}$"] = formatted
        else:
            replacements[f"${{{key}}}$"] = str(value) if pd.notna(value) else ""

    return replacements


def process_paragraph(paragraph, replacements):
    """Replace placeholders inside a single paragraph, preserving runs."""
    original_text = "".join(run.text for run in paragraph.runs)
    if not any(key in original_text for key in replacements):
        return

    for run in paragraph.runs:
        run.text = ""

    parts = re.split(
        f"({'|'.join(re.escape(k) for k in replacements)})", original_text
    )
    for part in parts:
        new_run = paragraph.add_run(
            part if part not in replacements else replacements[part]
        )
        if part in replacements:
            new_run.bold = True


def process_template(template_path, output_path, replacements):
    """Fill a DOCX template with replacements and save."""
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, replacements)

    doc.save(output_path)


# ══════════════════════════════════════════════
#  PHASE 1 — DOCX generation (fast, parallel)
# ══════════════════════════════════════════════

def generate_docx_for_consumer(idx, row, templates, template_dir, output_dir, dir_lock):
    """
    Generate all DOCX files for one consumer.
    Returns a list of (docx_path, pdf_path) pairs queued for later PDF conversion,
    plus the consumer directory path (needed for the post-merge step).
    """
    consumer_id   = str(row.get("CONSUMER_APP_NO",  f"consumer_{idx}"))
    consumer_name = str(row.get("CONSUMER_NAME",    "consumer"))
    safe_id        = sanitize_filename(consumer_id)
    safe_name      = sanitize_filename(consumer_name)

    with dir_lock:
        consumer_dir = Path(output_dir) / safe_name
        if consumer_dir.exists():
            print(f"  ↷  Skipping existing directory: {safe_name}")
            return safe_name, consumer_dir, []
        consumer_dir.mkdir(parents=True, exist_ok=True)

    replacements = create_replacements_dict(row)

    # Panel serial-number Excel side-output
    serial_str  = replacements.get("${PANEL_SR_NO}$", "")
    serial_list = [s.strip() for s in serial_str.split(",")]
    excel_path  = consumer_dir / f"{safe_name}_panel_serial_numbers.xlsx"
    pd.DataFrame(serial_list).to_excel(excel_path, index=False, header=False)
    print(f"  ✓  Excel created: {excel_path.name}")

    pending_pdfs = []   # (docx_path, pdf_path)

    for doc_type, template_file in templates.items():
        template_path = Path(template_dir) / template_file
        base          = f"{safe_name}_{doc_type}"
        docx_path     = consumer_dir / f"{base}.docx"
        pdf_path      = consumer_dir / f"{base}.pdf"

        try:
            process_template(template_path, docx_path, replacements)
            print(f"  ✓  DOCX [{doc_type}] → {safe_name}")
            pending_pdfs.append((docx_path, pdf_path))
        except Exception as e:
            print(f"  ✗  DOCX [{doc_type}] failed for {safe_name}: {e}")

    return safe_name, consumer_dir, pending_pdfs


# ══════════════════════════════════════════════
#  PHASE 2 — PDF conversion (slow, parallel)
# ══════════════════════════════════════════════

def _pdf_worker(args):
    """
    Multiprocessing worker — each process owns its own Word/COM instance.
    Must be a module-level function so multiprocessing can pickle it.
    """
    import pythoncom
    docx_path_str, pdf_path_str, max_size_mb = args
    docx_path = Path(docx_path_str)
    pdf_path  = Path(pdf_path_str)

    pythoncom.CoInitialize()
    try:
        convert(docx_path, pdf_path)
        pdf_size = os.path.getsize(pdf_path) / (1024 * 1024)
        if pdf_size > max_size_mb:
            return docx_path_str, pdf_path_str, False, f"exceeds {max_size_mb} MB ({pdf_size:.2f} MB)"
        return docx_path_str, pdf_path_str, True, None
    except Exception as e:
        return docx_path_str, pdf_path_str, False, str(e)
    finally:
        pythoncom.CoUninitialize()


def run_pdf_phase(all_pending, pdf_workers):
    """
    Convert all queued DOCX → PDF using multiprocessing.
    Each process gets its own isolated Word COM instance — no RPC conflicts.
    Returns a set of successfully created PDF path strings.
    """
    from multiprocessing import Pool

    print(f"\n── PDF conversion phase: {len(all_pending)} files, {pdf_workers} processes ──")
    success_pdfs = set()

    # Paths must be plain strings for cross-process pickling
    args_list = [
        (str(docx_path), str(pdf_path), MAX_PDF_SIZE_MB)
        for docx_path, pdf_path in all_pending
    ]

    with Pool(processes=pdf_workers) as pool:
        for docx_str, pdf_str, ok, err in pool.imap_unordered(_pdf_worker, args_list):
            name = Path(docx_str).name
            if ok:
                print(f"  ✓  PDF  ← {name}")
                success_pdfs.add(pdf_str)
            else:
                print(f"  ✗  PDF failed [{name}]: {err}")

    return success_pdfs


# ══════════════════════════════════════════════
#  PHASE 3 — PDF merge (per consumer, fast)
# ══════════════════════════════════════════════

def merge_consumer_pdfs(consumer_records):
    """
    Merge WCR + Aadhar PDFs for each consumer after all PDFs are ready.
    consumer_records: list of (safe_name, consumer_dir)
    """
    print(f"\n── PDF merge phase: {len(consumer_records)} consumers ──")
    for safe_name, consumer_dir in consumer_records:
        wcr_pdf      = consumer_dir / f"{safe_name}_WCR.pdf"
        aadhar_pdf   = consumer_dir / f"{safe_name}_Aadhar.pdf"
        combined_pdf = consumer_dir / f"{safe_name}_WCR_Aadhar_Combined.pdf"
 
        if wcr_pdf.exists() and aadhar_pdf.exists():
            if combine_pdf_files([wcr_pdf, aadhar_pdf], combined_pdf):
                print(f"  ✓  Combined PDF → {safe_name}")
            else:
                print(f"  ✗  Merge failed  → {safe_name}")
        else:
            missing = []
            if not wcr_pdf.exists():    missing.append("WCR")
            if not aadhar_pdf.exists(): missing.append("Aadhar")
            print(f"  ⚠  Missing {', '.join(missing)} PDF(s) for {safe_name}")


# ══════════════════════════════════════════════
#  ORCHESTRATOR
# ══════════════════════════════════════════════

def generate_documents(docx_workers=4, pdf_workers=6):
    """
    Two-phase document generation:
      Phase 1 — All DOCX files generated in parallel   (CPU-light, I/O-bound)
      Phase 2 — All PDF conversions run in parallel     (CPU-heavy, can use more threads)
      Phase 3 — WCR + Aadhar PDFs merged per consumer  (sequential, fast)

    Tune pdf_workers higher than docx_workers because PDF conversion is the
    bottleneck and benefits most from concurrency (up to your CPU core count).
    """
    INPUT_EXCEL  = "input_data.xlsx"
    TEMPLATE_DIR = "DOCS"
    OUTPUT_DIR   = "generated_docs"

    TEMPLATES = {
        "Annexure_1":           "TEMPELATE_Annexure.docx",
        "Aadhar":               "Aadhar.docx",
        "WCR":                  "WCR.docx",
        "Annexure_3":           "Annexure-3-Net-Metering.docx",
        "NP_Agreement":         "np_agreement.docx",
        "Meter_testing_letter": "Meter Testing Letter.docx",
        "DCR":                  "DCR.docx",
    }

    df = pd.read_excel(INPUT_EXCEL)
    Path(OUTPUT_DIR).mkdir(exist_ok=True)
    dir_lock = threading.Lock()

    total = len(df)
    print(f"Starting document generation for {total} consumers.")
    print(f"  DOCX workers : {docx_workers}")
    print(f"  PDF  workers : {pdf_workers}\n")

    # ── Phase 1: DOCX ──────────────────────────
    t0 = time.perf_counter()
    print("── DOCX generation phase ──")

    all_pending      = []   # [(docx_path, pdf_path), ...]
    consumer_records = []   # [(safe_name, consumer_dir), ...]

    with ThreadPoolExecutor(max_workers=docx_workers) as executor:
        futures = {
            executor.submit(
                generate_docx_for_consumer,
                idx, row, TEMPLATES, TEMPLATE_DIR, OUTPUT_DIR, dir_lock
            ): idx
            for idx, row in df.iterrows()
        }
        for future in as_completed(futures):
            try:
                safe_name, consumer_dir, pending = future.result()
                all_pending.extend(pending)
                consumer_records.append((safe_name, consumer_dir))
            except Exception as e:
                print(f"  ✗  DOCX phase error: {e}")

    t1 = time.perf_counter()
    print(f"\n  DOCX phase done in {t1 - t0:.1f}s — {len(all_pending)} files queued for PDF.\n")

    # ── Phase 2: PDF ───────────────────────────
    run_pdf_phase(all_pending, pdf_workers)

    t2 = time.perf_counter()
    print(f"\n  PDF phase done in {t2 - t1:.1f}s.\n")

    # ── Phase 3: Merge ─────────────────────────
    merge_consumer_pdfs(consumer_records)

    t3 = time.perf_counter()
    print(f"\n  Merge phase done in {t3 - t2:.1f}s.")
    print(f"\n✔  All done in {t3 - t0:.1f}s total.")


if __name__ == "__main__":
    # Recommended starting point:
    #   docx_workers=4  (scales with I/O, rarely needs more than 4)
    #   pdf_workers=6–8 (scale up to your CPU core count for best gains)
    generate_documents(docx_workers=4, pdf_workers=6)