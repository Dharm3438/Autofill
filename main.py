from docx import Document
import re
import pandas as pd

def replace_text_in_paragraph(paragraph, replacements):
    # Combine all run texts into a single string
    full_text = ''.join(run.text for run in paragraph.runs)

    # Create a pattern to match all replacement keys
    pattern = '|'.join(re.escape(k) for k in replacements.keys())
    parts = re.split(f'({pattern})', full_text)

    # Clear existing runs
    for run in paragraph.runs:
        run.text = ''

    # Rebuild paragraph with styled replacements
    for part in parts:
        run = paragraph.add_run()
        if part in replacements:
            run.text = replacements[part]
            run.bold = True
            run.underline = True
        else:
            run.text = part


def replace_text_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

def fill_template(template_path, output_path, replacements):
    doc = Document(template_path)

    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, replacements)

    doc.save(output_path)

# Example usage
replacements = {
    "${SOLAR_CAPACITY}$": "3",
    "${CONSUMER_NAME}$": "Indresh Devji Poladiya",
    "${CONSUMER_ADDRESS}$": "Ganesh Cross Road, Near Yogeshwar Medical Chalisgaon",
    "${CONSUMER_APP_NO}$": "123456789011",
    "${CONSUMER_APP_DATE}$": "01/01/2025",
    "${TOTAL_PANEL_CAPACITY}$": "3150 Watt",
    "${NO_OF_PANEL}$": "6",
    "${PANEL_SR_NO}$": "5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757,5167516757",
    "${PANEL_COMPANY}$": "Fujiyama Power System Private Limited",
    "${CELL_MANUFACTURER}$": "Mundra Solar Private Limited"
}

fill_template("DOCS\DCR.docx", "res_DCR.docx", replacements)

# # Read Excel file with the first row as headers (keys)
# df = pd.read_excel("input_data.xlsx", header=0, dtype=str)
# # Optional: print the keys to verify they are read correctly
# keys = df.columns.tolist()
# print("Excel Keys:", keys)

# # Iterate over each data row to build a replacements dictionary
# for index, row in df.iterrows():
#     # Build dictionary: wrap each key in ${} and map to its corresponding row value
#     replacements = {f"${{{col}}}$": row[col] for col in keys}
#     print(replacements)
#     # For naming the output file, use the value from the CONSUMER_NAME column
#     # consumer_name = row["CONSUMER_NAME"].replace(" ", "_")
#     output_filename = f"Dharmesh_DSR.docx"
    
#     # Call the fill_template function (assumed to be defined elsewhere)
#     fill_template("DOCS\DCR.docx", output_filename, replacements)
    
#     # Debug: print the dictionary to verify its contents
#     # print("Row dictionary:", replacements)

# # for index, row in df.iterrows():
# #     # Convert row data into a dictionary for replacements
# #     # replacements = {col: str(row[col]) for col in df.columns}
# #     # replacements = {f"${{{col}}}$": str(row[col]) for col in df.columns}
# #     # replacements = {f"${{{key}}}$": str(row[key]) for key in columns}
# #     # replacements = {f"${{{col}}}$": str(row[col]) for col in df.columns}
# #     replacements = {f"${{{keys[i]}}}$": str(row[i]) for i in range(len(keys))}

# #     print(replacements)
# #     # Generate filename using the CONSUMER_NAME column
# #     # consumer_name = row["#(CONSUMER_NAME)#"].replace(" ", "_")  # Replace spaces for safety
# #     consumer_name = "INDRESH_DEVJI_POLADIYA"
# #     output_filename = f"{consumer_name}_DSR.docx"

# #     # Call the function to fill the template
# #     fill_template("DOCS\DCR.docx", output_filename, replacements)

print("Documents generated successfully!")